"""Script para poblar el RAG con documentos existentes."""
import sys
import os

# Add backend to path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app.database import get_db
from app.models import DocumentoRAG
from docx import Document


def extract_text_from_docx(filepath):
    """Extrae texto de un archivo DOCX."""
    try:
        doc = Document(filepath)
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)

        return '\n'.join(paragraphs)
    except Exception as e:
        return f"[Error: {str(e)}]"


def populate_rag():
    """Poblar el RAG con documentos existentes."""
    db = next(get_db())

    # Documents to add
    documents = [
        {
            "titulo": "Políticas sobre Diseño Curricular UNAC",
            "tipo_documento": "politica",
            "categoria": "curricular",
            "fuente": "Docs/Viceacademica/Docs_ViceAcademica/2025.035.06.01 Políticas sobre Diseño Curricular UNAC.pdf",
            "contenido": "[PDF - requiere extracción manual o herramienta especializada]",
            "contenido_resumido": "Políticas institucionales sobre diseño curricular de la UNAC",
            "vigente": True,
            "etiquetas": ["diseño curricular", "políticas", "UNAC"],
            "metadatos": {"anio": 2025, "institucion": "UNAC"},
        },
        {
            "titulo": "Resolución 021795",
            "tipo_documento": "resolucion",
            "categoria": "normativo",
            "fuente": "Docs/Viceacademica/Docs_ViceAcademica/Resolución 021795.pdf",
            "contenido": "[PDF - requiere extracción manual o herramienta especializada]",
            "contenido_resumido": "Resolución institucional sobre procesos académicos",
            "vigente": True,
            "etiquetas": ["resolución", "normativa", "acreditación"],
            "metadatos": {"numero": "021795"},
        },
        {
            "titulo": "Inventario Documental UNAC Avanzado",
            "tipo_documento": "inventario",
            "categoria": "administrativo",
            "fuente": "Docs/Viceacademica/Docs_ViceAcademica/Inventario_Documental_UNAC_Avanzado.xlsx",
            "contenido": "[XLSX - requiere procesamiento especial]",
            "contenido_resumido": "Inventario de documentos institucionales UNAC",
            "vigente": True,
            "etiquetas": ["inventario", "documentos", "UNAC"],
        },
        {
            "titulo": "Formato C01 - Formato Oficial de Solicitud",
            "tipo_documento": "formato",
            "categoria": "calidad",
            "fuente": "Docs/Viceacademica/Calidad/compartoformatosrc/Formato_oficial_C01.docx",
            "contenido": "[Pendiente de extracción]",
            "contenido_resumido": "Formato oficial C01 para procesos de calidad",
            "vigente": True,
            "etiquetas": ["formato", "calidad", "C01"],
        },
        {
            "titulo": "Reunión 23-04-2026 - Vicerrectoría Académica",
            "tipo_documento": "acta",
            "categoria": "administrativo",
            "fuente": "Docs/Viceacademica/Requisitos_Viceacademico/Reunión 23-04-2026.docx",
            "contenido": "[Pendiente de extracción]",
            "contenido_resumido": "Acta de reunión de Vicerrectoría Académica",
            "vigente": True,
            "etiquetas": ["acta", "vicerrectoría", "reunión"],
            "metadatos": {"fecha": "2026-04-23"},
        },
    ]

    # Try to extract actual content from DOCX files
    base_path = "E:/Projects/appsgeap/Docs/Viceacademica/Calidad/compartoformatosrc"

    format_files = [
        ("Formato_oficial_C01.docx", "Formato C01"),
        ("Formato_oficial_C02.docx", "Formato C02"),
        ("Formato_oficial_C03.docx", "Formato C03"),
        ("Formato_oficial_C04.docx", "Formato C04"),
        ("Formato_oficial_C05.docx", "Formato C05"),
        ("Formato_oficial_C06.docx", "Formato C06"),
        ("Formato_oficial_C07.docx", "Formato C07"),
        ("Formato_oficial_C08.docx", "Formato C08"),
    ]

    for filename, label in format_files:
        filepath = f"{base_path}/{filename}"
        if os.path.exists(filepath):
            contenido = extract_text_from_docx(filepath)
            if len(contenido) > 100:
                documents.append({
                    "titulo": f"Formato {label} - {filename}",
                    "tipo_documento": "formato",
                    "categoria": "calidad",
                    "fuente": filepath,
                    "contenido": contenido,
                    "contenido_resumido": contenido[:500],
                    "vigente": True,
                    "etiquetas": ["formato", "calidad", label.lower()],
                })

    # Try to extract from the meeting document
    meeting_path = "E:/Projects/appsgeap/Docs/Viceacademica/Requisitos_Viceacademico/Reunión 23-04-2026.docx"
    if os.path.exists(meeting_path):
        contenido = extract_text_from_docx(meeting_path)
        # Find and update the meeting document
        for doc in documents:
            if "Reunión" in doc["titulo"]:
                doc["contenido"] = contenido
                doc["contenido_resumido"] = contenido[:500]
                break

    # Add to database
    created = 0
    for doc_data in documents:
        # Check if already exists
        existing = db.query(DocumentoRAG).filter(
            DocumentoRAG.titulo == doc_data["titulo"]
        ).first()

        if not existing:
            # Remove None values to let SQLAlchemy handle defaults
            doc_data_clean = {k: v for k, v in doc_data.items() if v is not None}
            if "creado_por" in doc_data_clean:
                del doc_data_clean["creado_por"]

            doc = DocumentoRAG(**doc_data_clean)
            db.add(doc)
            created += 1

    db.commit()
    print(f"Documentos RAG creados: {created}")
    print(f"Total de documentos en RAG: {db.query(DocumentoRAG).count()}")

    # List all documents
    print("\nDocumentos en RAG:")
    for doc in db.query(DocumentoRAG).order_by(DocumentoRAG.fecha_creacion.desc()).all():
        print(f"  - {doc.titulo[:60]}... ({doc.categoria})")


if __name__ == "__main__":
    populate_rag()