"""RAG institutional router with pgvector semantic search."""
import unicodedata
import re
import time
from typing import Optional
from uuid import UUID

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Query
from sqlalchemy.orm import Session
from sqlalchemy import func, or_, text

from app.database import get_db
from app.models import DocumentoRAG
from app.schemas.documento_rag import (
    DocumentoRAGCreate,
    DocumentoRAGResponse,
    BusquedaRAGRequest,
    BusquedaRAGResponse,
    ResultadoBusquedaRAG,
)
from app.models import User
from app.services.embedding import generate_embedding, vector_to_sql

from sqlalchemy import create_engine

RAG_DB_URL = "postgresql+pg8000://rag_user:rag_pass_2026@localhost:5434/sgeap_rag"
rag_engine = create_engine(RAG_DB_URL)

router = APIRouter(prefix="/api/rag", tags=["rag"])


def normalize_text(text: str) -> str:
    text = text.lower().strip()
    text = re.sub(r'\s+', ' ', text)
    return text


def extract_text_from_docx(content: bytes) -> str:
    try:
        from docx import Document
        import io
        doc = Document(io.BytesIO(content))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)
        return '\n'.join(paragraphs)
    except Exception as e:
        return f"[Error extracting DOCX: {str(e)}]"


def extract_text_from_pdf(content: bytes) -> str:
    try:
        import io
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return '\n'.join(text_parts)
        except ImportError:
            pass
        try:
            from PyPDF2 import PdfReader
            text_parts = []
            reader = PdfReader(io.BytesIO(content))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return '\n'.join(text_parts)
        except ImportError:
            pass
        return "[PDF extraction not available]"
    except Exception as e:
        return f"[Error extracting PDF: {str(e)}]"


def calculate_relevance(query: str, content: str, titulo: str) -> float:
    query_words = set(normalize_text(query).split())
    content_lower = normalize_text(content)
    titulo_lower = normalize_text(titulo)
    score = 0.0
    for word in query_words:
        if word in titulo_lower:
            score += 0.4
        if word in content_lower:
            score += 0.1
    if query_words:
        score = min(score / len(query_words) * 2, 1.0)
    return round(score, 3)


def get_snippet(content: str, query: str, max_length: int = 300) -> str:
    content_lower = normalize_text(content)
    query_lower = normalize_text(query)
    query_words = query_lower.split()
    first_pos = len(content)
    for word in query_words:
        pos = content_lower.find(word)
        if pos != -1 and pos < first_pos:
            first_pos = pos
    if first_pos == len(content):
        return content[:max_length] + "..."
    start = max(0, first_pos - 50)
    end = min(len(content), first_pos + max_length)
    snippet = content[start:end].strip()
    if start > 0:
        snippet = "..." + snippet
    if end < len(content):
        snippet = snippet + "..."
    return snippet


@router.get("/documentos", response_model=list[DocumentoRAGResponse])
def list_documentos(
    skip: int = 0,
    limit: int = 100,
    categoria: Optional[str] = None,
    db: Session = Depends(get_db),
):
    """List all RAG documents."""
    query = db.query(DocumentoRAG)
    if categoria:
        query = query.filter(DocumentoRAG.categoria == categoria)
    documentos = query.order_by(DocumentoRAG.fecha_creacion.desc()).offset(skip).limit(limit).all()
    return [DocumentoRAGResponse.model_validate(d) for d in documentos]


@router.get("/documentos/{documento_id}", response_model=DocumentoRAGResponse)
def get_documento(documento_id: UUID, db: Session = Depends(get_db)):
    """Get a specific RAG document."""
    documento = db.query(DocumentoRAG).filter(DocumentoRAG.id == documento_id).first()
    if not documento:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
    return DocumentoRAGResponse.model_validate(documento)


@router.post("/documentos", response_model=DocumentoRAGResponse, status_code=201)
def create_documento(
    documento_data: DocumentoRAGCreate,
    db: Session = Depends(get_db),
):
    """Create a new RAG document."""
    documento = DocumentoRAG(
        titulo=documento_data.titulo,
        tipo_documento=documento_data.tipo_documento,
        categoria=documento_data.categoria,
        fuente=documento_data.fuente,
        url_origen=documento_data.url_origen,
        contenido=documento_data.contenido,
        contenido_resumido=documento_data.contenido_resumido,
        metadatos=documento_data.metadatos,
        etiquetas=documento_data.etiquetas,
        vigente=documento_data.vigente,
        fecha_documento=documento_data.fecha_documento,
        search_vector=normalize_text(documento_data.contenido),
    )
    db.add(documento)
    db.commit()
    db.refresh(documento)
    return DocumentoRAGResponse.model_validate(documento)


@router.post("/documentos/upload")
async def upload_documento(
    file: UploadFile = File(...),
    titulo: str = Query(...),
    tipo_documento: str = Query(...),
    categoria: str = Query(...),
    db: Session = Depends(get_db),
):
    """Upload and extract text from a document (DOCX or PDF) with semantic embeddings."""
    content = await file.read()
    if file.filename.endswith('.docx'):
        contenido = extract_text_from_docx(content)
    elif file.filename.endswith('.pdf'):
        contenido = extract_text_from_pdf(content)
    else:
        raise HTTPException(status_code=400, detail="Solo se permiten archivos DOCX o PDF")

    if len(contenido) < 50:
        raise HTTPException(status_code=400, detail="No se pudo extraer texto del documento")

    documento = DocumentoRAG(
        titulo=titulo,
        tipo_documento=tipo_documento,
        categoria=categoria,
        fuente=file.filename,
        contenido=contenido,
        contenido_resumido=contenido[:500] if len(contenido) > 500 else contenido,
        search_vector=normalize_text(contenido),
    )
    db.add(documento)
    db.commit()
    db.refresh(documento)

    try:
        embedding = generate_embedding(contenido[:8000])
        vector_str = vector_to_sql(embedding)
        with rag_engine.connect() as conn:
            conn.execute(text("""
                INSERT INTO documentos_rag_embeddings
                (id, titulo, tipo_documento, categoria, fuente, contenido, contenido_resumido, embedding)
                VALUES (:id, :titulo, :tipo, :categoria, :fuente, :contenido, :resumido, CAST(:embedding AS vector))
            """), {
                "id": str(documento.id),
                "titulo": titulo,
                "tipo": tipo_documento,
                "categoria": categoria,
                "fuente": file.filename,
                "contenido": contenido,
                "resumido": contenido[:500] if len(contenido) > 500 else contenido,
                "embedding": vector_str,
            })
            conn.commit()
    except Exception as e:
        print(f"Warning: Could not create embedding: {e}")

    return {
        "id": str(documento.id),
        "titulo": documento.titulo,
        "extracto": contenido[:200] + "...",
        "total_caracteres": len(contenido),
    }


@router.post("/buscar", response_model=BusquedaRAGResponse)
def buscar_documentos(
    busqueda: BusquedaRAGRequest,
    db: Session = Depends(get_db),
):
    """Search RAG documents based on query using semantic similarity."""
    start_time = time.time()
    query_embedding = generate_embedding(busqueda.query)
    query_vector = vector_to_sql(query_embedding)

    try:
        with rag_engine.connect() as conn:
            sql = """
                SELECT id, titulo, tipo_documento, categoria, fuente, contenido_resumido, embedding,
                       1 - (embedding <=> CAST(:query_vector AS vector)) as similarity
                FROM documentos_rag_embeddings
                WHERE vigente = TRUE
            """
            params = {"query_vector": query_vector}
            if busqueda.categoria:
                sql += " AND categoria = :categoria"
                params["categoria"] = busqueda.categoria
            if busqueda.tipo_documento:
                sql += " AND tipo_documento = :tipo"
                params["tipo"] = busqueda.tipo_documento
            sql += " ORDER BY embedding <=> CAST(:query_vector AS vector) LIMIT :limite"
            params["limite"] = busqueda.limite

            result = conn.execute(text(sql), params)
            rows = result.fetchall()

            resultados = []
            for row in rows:
                snippet = get_snippet(row.contenido_resumido or "", busqueda.query, 300)
                resultados.append(ResultadoBusquedaRAG(
                    id=row.id,
                    titulo=row.titulo,
                    tipo_documento=row.tipo_documento,
                    categoria=row.categoria,
                    fuente=row.fuente,
                    contenido_resumido=row.contenido_resumido,
                    snippet=snippet,
                    relevancia=round(row.similarity, 3),
                    url_origen=None,
                ))

            elapsed_ms = int((time.time() - start_time) * 1000)
            return BusquedaRAGResponse(
                resultados=resultados,
                total=len(resultados),
                query=busqueda.query,
                tiempo_ms=elapsed_ms,
            )
    except Exception as e:
        print(f"Semantic search error: {e}")

    query = db.query(DocumentoRAG)
    if busqueda.categoria:
        query = query.filter(DocumentoRAG.categoria == busqueda.categoria)
    if busqueda.tipo_documento:
        query = query.filter(DocumentoRAG.tipo_documento == busqueda.tipo_documento)

    documentos = query.all()
    resultados = []
    for doc in documentos:
        score = calculate_relevance(busqueda.query, doc.contenido, doc.titulo)
        if score >= 0.1:
            snippet = get_snippet(doc.contenido, busqueda.query)
            resultados.append(ResultadoBusquedaRAG(
                id=doc.id,
                titulo=doc.titulo,
                tipo_documento=doc.tipo_documento,
                categoria=doc.categoria,
                fuente=doc.fuente,
                contenido_resumido=doc.contenido_resumido,
                snippet=snippet,
                relevancia=score,
                url_origen=doc.url_origen,
            ))

    resultados.sort(key=lambda x: x.relevancia, reverse=True)
    resultados = resultados[:busqueda.limite]

    elapsed_ms = int((time.time() - start_time) * 1000)
    return BusquedaRAGResponse(
        resultados=resultados,
        total=len(resultados),
        query=busqueda.query,
        tiempo_ms=elapsed_ms,
    )


@router.get("/sugerencias")
def get_sugerencias_fase(
    fase_nombre: str = Query(...),
    programa_nombre: str = Query(default=""),
    db: Session = Depends(get_db),
):
    """Get RAG suggestions for filling a phase."""
    query = f"{fase_nombre}"
    if programa_nombre:
        query += f" {programa_nombre}"

    start_time = time.time()
    try:
        query_embedding = generate_embedding(query)
        query_vector = vector_to_sql(query_embedding)

        with rag_engine.connect() as conn:
            sql = """
                SELECT id, titulo, tipo_documento, categoria, fuente, contenido_resumido,
                       1 - (embedding <=> CAST(:query_vector AS vector)) as similarity
                FROM documentos_rag_embeddings
                WHERE vigente = TRUE
                ORDER BY embedding <=> CAST(:query_vector AS vector)
                LIMIT 5
            """
            result = conn.execute(text(sql), {"query_vector": query_vector})
            rows = result.fetchall()

            resultados = []
            for row in rows:
                if row.similarity >= 0.3:
                    resultados.append({
                        "id": row.id,
                        "titulo": row.titulo,
                        "tipo_documento": row.tipo_documento,
                        "categoria": row.categoria,
                        "fuente": row.fuente,
                        "contenido_resumido": row.contenido_resumido,
                        "relevancia": round(row.similarity, 3),
                    })

            return {
                "sugerencias": resultados,
                "total": len(resultados),
                "fase": fase_nombre,
                "programa": programa_nombre,
                "tiempo_ms": int((time.time() - start_time) * 1000),
            }
    except Exception as e:
        print(f"Suggestions error: {e}")

    docs = db.query(DocumentoRAG).filter(
        or_(
            DocumentoRAG.titulo.ilike(f"%{fase_nombre}%"),
            DocumentoRAG.contenido.ilike(f"%{fase_nombre}%"),
        )
    ).limit(5).all()

    sugerencias = []
    for doc in docs:
        score = calculate_relevance(fase_nombre, doc.contenido, doc.titulo)
        if score >= 0.1:
            sugerencias.append({
                "id": doc.id,
                "titulo": doc.titulo,
                "tipo_documento": doc.tipo_documento,
                "categoria": doc.categoria,
                "fuente": doc.fuente,
                "contenido_resumido": doc.contenido_resumido,
                "relevancia": score,
            })

    return {
        "sugerencias": sugerencias,
        "total": len(sugerencias),
        "fase": fase_nombre,
        "programa": programa_nombre,
        "tiempo_ms": int((time.time() - start_time) * 1000),
    }


@router.delete("/documentos/{documento_id}", status_code=204)
def delete_documento(documento_id: UUID, db: Session = Depends(get_db)):
    """Delete a RAG document."""
    documento = db.query(DocumentoRAG).filter(DocumentoRAG.id == documento_id).first()
    if not documento:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
    db.delete(documento)
    db.commit()
    return None


@router.get("/documentos/{documento_id}/contenido")
def get_documento_contenido(documento_id: UUID, db: Session = Depends(get_db)):
    """Get full content of a document for highlighting."""
    rag_engine_local = create_engine(RAG_DB_URL)
    try:
        with rag_engine_local.connect() as conn:
            result = conn.execute(
                text("""
                    SELECT contenido, titulo, tipo_documento, categoria
                    FROM documentos_rag_embeddings
                    WHERE id = :id
                """),
                {"id": str(documento_id)}
            )
            row = result.fetchone()
            if row:
                return {
                    "contenido": row[0],
                    "titulo": row[1],
                    "tipo_documento": row[2],
                    "categoria": row[3],
                }
    except Exception as e:
        print(f"RAG lookup error: {e}")

    raise HTTPException(status_code=404, detail="Contenido no encontrado")


@router.get("/documentos/{documento_id}/archivo")
def get_documento_archivo(documento_id: UUID, db: Session = Depends(get_db)):
    """Get the PDF file path for a document."""
    from pathlib import Path
    from fastapi.responses import FileResponse

    def normalize(s: str) -> str:
        if not s:
            return ""
        return unicodedata.normalize('NFD', str(s)).encode('ascii', 'ignore').decode('ascii', errors='ignore')

    docs_base = Path("E:/Projects/appsgeap/Docs")
    all_pdfs = {normalize(p.name.lower()): p for p in docs_base.rglob("*.pdf")}
    print(f"[RAG] Looking for doc ID: {documento_id}, found {len(all_pdfs)} PDFs")

    rag_engine_local = create_engine(RAG_DB_URL)
    try:
        with rag_engine_local.connect() as conn:
            result = conn.execute(
                text("SELECT titulo, fuente FROM documentos_rag_embeddings WHERE id = :id"),
                {"id": str(documento_id)}
            )
            row = result.fetchone()
            if row:
                titulo = normalize(row[0])
                fuente = normalize(row[1])
                print(f"[RAG] Query result - titulo: {row[0]}, fuente: {row[1]}")

                keywords = [w for w in titulo.split() if len(w) > 3]
                print(f"[RAG] Keywords: {keywords}")

                for pdf_name_lower, pdf_path in all_pdfs.items():
                    score = sum(1 for kw in keywords if kw.lower() in pdf_name_lower)
                    if score >= 2:
                        print(f"[RAG] Match found: {pdf_path.name} (score={score})")
                        return FileResponse(
                            str(pdf_path),
                            media_type="application/pdf",
                            filename=pdf_path.name,
                            headers={"Content-Disposition": "inline; filename*=UTF-8''" + pdf_path.name}
                        )
                print(f"[RAG] No PDF match for keywords: {keywords}")

    except Exception as e:
        print(f"[RAG] Error: {e}")

    raise HTTPException(status_code=404, detail="Archivo PDF no encontrado")


@router.get("/categorias")
def list_categorias(db: Session = Depends(get_db)):
    """List all document categories with counts."""
    results = db.query(
        DocumentoRAG.categoria,
        func.count(DocumentoRAG.id).label('count')
    ).group_by(DocumentoRAG.categoria).all()
    return [{"categoria": r.categoria, "count": r.count} for r in results]


@router.get("/tipos")
def list_tipos(db: Session = Depends(get_db)):
    """List all document types."""
    results = db.query(
        DocumentoRAG.tipo_documento,
        func.count(DocumentoRAG.id).label('count')
    ).group_by(DocumentoRAG.tipo_documento).all()
    return [{"tipo": r.tipo_documento, "count": r.count} for r in results]